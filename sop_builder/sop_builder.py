import requests  # for perplexity
from fpdf import FPDF  # to download sop as pdf
from docx import Document  # to download sop as doc
import re
import os
import PyPDF2  # for cv upload and parsing

# === Extraction helpers ===

def extract_name_from_cv(text):
    lines = text.strip().split('\n')
    lines = [line.strip() for line in lines if line.strip()]
    # strategy 1: Look for title-style capitalized name line on top
    for line in lines[:5]:
        if (
            re.match(r"^[A-Z][a-z]+(?: [A-Z][a-z]+)+$", line) and
            len(line.split()) <= 4 and
            not any(char in line for char in ['|', '@', 'http', '/', '\\', ':'])
        ):
            return line
    # strategy 2: Try "Name: John Doe"
    match = re.search(r'Name[:\-]\s*(.+)', text, re.IGNORECASE)
    if match:
        possible_name = match.group(1).strip()
        if len(possible_name.split()) <= 4:
            return possible_name
    return "Your Name"

def extract_academic_qualifications(text):
    match = re.search(r'(EDUCATION|ACADEMIC QUALIFICATIONS)(.*?)(PROJECTS|SKILLS|EXPERIENCE|ACHIEVEMENTS|$)', text, re.IGNORECASE | re.DOTALL)
    if match:
        degrees = match.group(2).strip()
        degrees = re.sub(r'\n+', '\n', degrees)
        return degrees
    return ""

def determine_intended_degree(academic_text):
    academic_text = academic_text.lower()
    if "phd" in academic_text:
        return "PhD"
    elif "master" in academic_text or "msc" in academic_text:
        return "PhD"
    elif "bachelor" in academic_text:
        return "Masters"
    else:
        return "Masters"

def extract_text_from_docx(filepath):
    try:
        doc = Document(filepath)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception:
        return ""

def extract_text_from_pdf(filepath):
    try:
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
    except Exception:
        return ""

def extract_section(text, headings):
    pattern = '|'.join(headings)
    match = re.search(rf'(?i)({pattern})[:\s]*([\s\S]+?)(?:\n\s*\n|$)', text)
    return match.group(2).strip() if match else None

def parse_cv(text):
    return {
        "name": extract_name_from_cv(text),
        "academic_qualifications": extract_academic_qualifications(text),
        "intended_degree": determine_intended_degree(extract_academic_qualifications(text)),
        "key_skills": extract_section(text, ["skills", "technical skills"]),
        "projects": extract_section(text, ["projects", "publications", "research"]),
        "awards": extract_section(text, ["awards", "scholarships", "recognitions"]),
        "hobbies": extract_section(text, ["hobbies", "volunteer work", "extracurriculars"])
    }

# === SOP prompt builder ===

def build_sop_prompt(user_inputs):
    name = user_inputs.get("name")
    country_of_origin = user_inputs.get("country_of_origin")
    intended_degree = user_inputs.get("intended_degree")
    preferred_country = user_inputs.get("preferred_country")
    field_of_study = user_inputs.get("field_of_study")
    preferred_uni = user_inputs.get("preferred_uni")

    base_prompt = (
        f"I am {name}, I am from {country_of_origin}. I want to study {intended_degree} in {preferred_country}. "
        f"My preferred field of study is {field_of_study}. "
        f"My preferred university is {preferred_uni}. "
        f"I want you to write me a SOP"
    )

    base_prompt += (
        "Make sure the SOP is ATS friendly and does not look like an AI wrote it. "
        "It should look like a human wrote it. "
        "The SOP should follow a professional and formal tone."
        "Only respond with the SOP text. Do not include any explanations or additional messages."
        "Exclude all inline citations or footnote markers.\n"
    )

    base_prompt += "Here are my details:\n"

    optional_fields = [
        ("key_skills", "My key skills are"),
        ("strengths", "My strengths are"),
        ("why_field", "I want to pursue this field because"),
        ("why_uni", "And this university because"),
        ("goals", "My long term goals are"),
        ("challenge", "More about me:")
    ]

    for key, label in optional_fields:
        value = user_inputs.get(key)
        if value:
            base_prompt += f"{label} {value}.\n"

    if len(user_inputs.get("projects")) > 0:
        base_prompt += "I have also completed these projects/research/publications:\n"
        for p in user_inputs.get("projects"):
            base_prompt += (f"Type: {p.get("type")}\n"
                            f"Title: {p.get("title")}\n"
                            f"Link: {p.get("link")}\n"         
                            f"Description: {p.get("description")}\n\n"   
                            )
            
    if len(user_inputs.get("education")) > 0:
        base_prompt += "Here are my past education details:\n"
        for e in user_inputs.get("education"):
            if e.get("isPresent"): end = "Presently studying here"
            else: end = e.get("endDate")
            if e.get("universityName") == "Other": uni = e.get("otherUniversityName")
            else: uni = e.get("universityName")
            base_prompt += (f"Discipline: {e.get("discipline")}\n"
                            f"Course Name: {e.get("course")}\n"         
                            f"Level of Study: {e.get("level")}\n"   
                            f"Country of Study: {e.get("country")}\n"   
                            f"Location of Study: {e.get("location")}\n"
                            f"Results I got: {e.get("results")}\n"
                            f"University Name: {uni}\n"
                            f"Start Date: {e.get("startDate")}\n"
                            f"End Date: {end}\n"
                            )
            
    if len(user_inputs.get("awards")) > 0:
        base_prompt += "I have received these certifications:\n"
        for p in user_inputs.get("awards"):
            base_prompt += (f"Type of certification: {p.get("type")}\n"
                            f"Name: {p.get("name")}\n"
                            f"Issuing Organization: {p.get("organization")}\n"         
                            f"Date obtained: {p.get("dateObtained")}\n\n"   
                            )
            
    if len(user_inputs.get("activity")) > 0:
        base_prompt += "Here is more information about the activities I did:\n"
        for p in user_inputs.get("activity"):
            base_prompt += (f"{p.get("type")}: {p.get("description")}")
        
    return base_prompt.strip()

def call_perplexity_api(prompt, token):
    url = "https://api.perplexity.ai/chat/completions"
    payload = {
        "model": "sonar",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 2048
    }
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    response = requests.post(url, json=payload, headers=headers)
    return response.json()

# === Exported function for API use ===

def generate_sop(user_inputs, token):
    """CORE function for Flask API, returns (sop, prompt)"""
    prompt = build_sop_prompt(user_inputs)
    response = call_perplexity_api(prompt, token)
    sop = response.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
    return sop, prompt

# === PDF/DOCX helpers (not usually in API, but retained for CLI or download endpoint) ===

def clean_text_for_pdf(text):
    text = text.replace("—", "-").replace("–", "-")
    text = text.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")
    return text.encode("latin-1", "ignore").decode("latin-1")

def save_pdf(filename, content):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    content = clean_text_for_pdf(content)
    pdf.multi_cell(0, 10, content)
    pdf.output(filename)

def save_docx(filename, content):
    doc = Document()
    doc.add_paragraph(content)
    doc.save(filename)
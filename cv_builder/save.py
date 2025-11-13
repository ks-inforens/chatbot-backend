import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def ensure_full_url(url):
    if not url.startswith(("http://", "https://")):
        return "https://" + url
    return url

def add_hyperlink(paragraph, url, text, color=RGBColor(0, 0, 255), underline=True):
    part = paragraph.part
    r_id = part.relate_to(ensure_full_url(url), "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), str(color))

    rPr.append(c)

    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return None

def add_bottom_border(paragraph):
    """Add a horizontal line under section headers"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")  # thicker line
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

def normalize_text(text):
    if text.strip().startswith("```"):
        text = text.strip().split("```")[1]
        if text.lower().startswith("json"):
            text = text[4:].strip()
    return text.strip()

def save_as_docx(text, filename="generated_cv.docx"):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    try:
        data = json.loads(text)
        is_json = True
    except:
        is_json = False

    if not is_json:
        for line in text.split("\n"):
            para = doc.add_paragraph(line.strip())
            para.paragraph_format.space_after = Pt(4)
        doc.save(filename)
        print(f"DOCX saved as {filename}")
        return filename

    # Helper function for section headers
    def add_section_header(title):
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        add_bottom_border(para)
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(8)

    # === Full Name and Location ===
    full_name = data.get("full_name", "")
    if full_name:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(full_name)
        run.bold = True
        run.font.size = Pt(16)

    location = data.get("location", "")
    if location:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(location)
        run.font.size = Pt(10)

    # === Contact Information ===
    email = data.get("email", "")
    phone = data.get("phone", "")
    links = data.get("links", [])

    if email or phone or links:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        first = True
        if email:
            add_hyperlink(para, f"mailto:{email}", email)
            first = False
        if phone:
            if not first:
                para.add_run(" | ")
            para.add_run(phone)
            first = False
        for link in links:
            if not first:
                para.add_run(" | ")
            url = link.get("url", "")
            name = link.get("name", url)
            if url:
                add_hyperlink(para, url, url)
                first = False

    # === 1. Professional Statement ===
    prof_stmt = data.get("professional_statement", "")
    if prof_stmt:
        add_section_header("Professional Statement")
        doc.add_paragraph(prof_stmt)

    # Prepare conditional order for Work Experience and Education
    work_exp = data.get("work_experience", [])
    education = data.get("education", [])

    # === 2 & 3. Work Experience or Education based on work experience count ===
    def write_work_experience():
        if work_exp:
            add_section_header("Work Experience")
            for job in work_exp:
                job_title = job.get("job_title", "")
                company_name = job.get("company_name", "")
                start_date = job.get("start_date", "")
                end_date = job.get("end_date", "")
                dates = ""
                if start_date and end_date:
                    dates = f"{start_date} - {end_date}"
                elif start_date:
                    dates = f"{start_date} - Present"

                para = doc.add_paragraph()
                run = para.add_run(job_title)
                run.bold = True
                run.italic = True
                if company_name:
                    run_comp = para.add_run(f" | {company_name}")
                    run_comp.bold = True
                    run_comp.italic = True
                if dates:
                    tab_stop = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
                    para.paragraph_format.tab_stops.add_tab_stop(tab_stop, alignment=WD_TAB_ALIGNMENT.RIGHT)
                    run_date = para.add_run(f"\t{dates}")
                    run_date.italic = True
                
                for resp in job.get("responsibilities", []):
                    bullet = doc.add_paragraph(resp, style='List Bullet')
                    bullet.paragraph_format.space_after = Pt(2)

                for ach in job.get("achievements", []):
                    bullet = doc.add_paragraph("Achievement: " + ach, style='List Bullet')
                    bullet.paragraph_format.space_after = Pt(2)

    def write_education():
        if education:
            add_section_header("Education")
            for edu in education:
                uni = edu.get("university_name", "")
                course = edu.get("course", "")
                discipline = edu.get("discipline", "")
                result = edu.get("results", "")
                start_date = edu.get("start_date", "")
                end_date = edu.get("end_date", "")
                dates = ""
                if start_date and end_date:
                    dates = f"{start_date} - {end_date}"
                elif start_date:
                    dates = f"{start_date} - Present"

                para = doc.add_paragraph()
                run = para.add_run(uni)
                run.bold = True
                run.italic = True
                if dates:
                    tab_stop = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
                    para.paragraph_format.tab_stops.add_tab_stop(tab_stop, alignment=WD_TAB_ALIGNMENT.RIGHT)
                    run_date = para.add_run(f"\t{dates}")
                    run_date.italic = True

                degree_field = f"{course} - {discipline}".strip(" -")
                if degree_field:
                    doc.add_paragraph(degree_field)
                if result:
                    doc.add_paragraph(f"Result: {result}")

    if len(work_exp) > 1:
        # Work Experience is 2nd, Education is 3rd
        write_work_experience()
        write_education()
    else:
        # Education is 2nd, Work Experience is 3rd
        write_education()
        write_work_experience()

    # === 4. Projects, Publications or Research ===
    projects = data.get("projects", [])
    if projects:
        add_section_header("Projects and Publications")
        for proj in projects:
            title = proj.get("title", "")
            desc = proj.get("description", "")
            type_ = proj.get("type", "")

            para = doc.add_paragraph()
            run = para.add_run(title)
            run.bold = True
            run.italic = True
            if type_:
                run_type = para.add_run(f" [{type_}]")
                run_type.italic = True
            
            if desc:
                doc.add_paragraph(desc, style='List Bullet')

    # === 5. Skills ===
    skills = data.get("skills", [])
    if skills:
        add_section_header("Skills")
        doc.add_paragraph(", ".join(skills))

    # === 6. Certifications ===
    certs = data.get("certifications", [])
    if certs:
        add_section_header("Certifications and Awards")
        for cert in certs:
            name = cert.get("name", "")
            organisation = cert.get("organisation", "")
            date = cert.get("date", "")
            type_ = cert.get("type", "")

            line = f"{name}"
            if organisation:
                line += f", {organisation}"
            if date:
                line += f" ({date})"
            if type_:
                line += f" [{type_}]"

            doc.add_paragraph(line, style='List Bullet')

    # === 7. Languages ===
    languages = data.get("languages_known", [])
    if languages:
        add_section_header("Languages Known")
        doc.add_paragraph(", ".join(languages))

    # === 8. Additional Sections ===
    additional_sections = data.get("additionalSec", [])
    for add_sec in additional_sections:
        title = add_sec.get("title", "")
        desc = add_sec.get("desc", "")
        if any(word in title.lower() for word in ["professional", "statement", "summary"]):
            pass
        elif title and desc:
            add_section_header(title)
            for paragraph in desc.split("\n"):
                doc.add_paragraph(paragraph.strip())

    doc.save(filename)
    print(f"DOCX saved as {filename}")
    return filename